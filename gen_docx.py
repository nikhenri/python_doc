# -------------------------------------------------------------------------------------
#  Description:
# -------------------------------------------------------------------------------------
#  Author: Nik Henri
# -------------------------------------------------------------------------------------
import os
import sys
import subprocess
import re
import zipfile

try:
    import docx
except ImportError:
    print("Installing 'docx' and 'docxtpl'")
    subprocess.check_call([sys.executable, "-m", "pip", "install", 'docx'])
    subprocess.check_call([sys.executable, "-m", "pip", "install", 'docxtpl'])
finally:
    import docx

from docx.shared import Cm
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.oxml import OxmlElement
import lxml


# -------------------------------------------------------------------------------------
class GenDocx:
    DOC_PARSER_CAPTION = 1
    MARK_ENTRY_COLOR = docx.shared.RGBColor(255, 192, 0)

    # -------------------------------------------------------------------------------------
    def __init__(self, register_dict):
        self.register_dict = register_dict
        self.document = None

    # -------------------------------------------------------------------------------------
    def generate(self, filename, open_generate_file=False):
        self.document = docx.Document()
        sections = self.document.sections
        for section in sections:
            section.left_margin = Cm(2.0)

        self.add_module_table(filename)
        for module in self.register_dict.keys():
            p = self.document.add_heading(module, level=0)
            self.add_bookmark(paragraph=p, bookmark_text="", bookmark_name=module)
            self.add_ip_docmentation(module)
            self.document.add_page_break()

        self.document.save(f'{filename}.docx')
        # # self.set_updatefields_true(f'{filename}.docx')

        with zipfile.ZipFile(f'{filename}.docx') as z:
            with open(f'{filename}.xml', 'wb') as f:
                f.write(z.read('word/document.xml'))

        if open_generate_file:
            os.startfile(f'{filename}.docx')

    # -------------------------------------------------------------------------------------
    def add_module_table(self, filename):
        self.document.add_heading(filename.upper() + " registers", level=0)
        table = self.document.add_table(rows=1, cols=1)
        shading_elm = []
        for i in range(len(table.columns)):  # put the background gray for the 1er row
            shading_elm.append(docx.oxml.parse_xml(r'<w:shd {} w:fill="d9d9d9"/>'.format(docx.oxml.ns.nsdecls('w'))))
            table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading_elm[i])

        table.style = 'Table Grid'  # add border
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Modules'
        #hdr_cells[2].width = Cm(19.0)

        for module in self.register_dict.keys():
            row_cells = table.add_row().cells
            self.add_link(paragraph=row_cells[0].paragraphs[0], link_to=module, text=module, tool_tip=module)
            #row_cells[0].width = Cm(19.0)
        self.document.add_page_break()

    # -------------------------------------------------------------------------------------
    def add_ip_docmentation(self, module):
        self.add_summary_table(module)
        self.add_register_table(module)

    # -------------------------------------------------------------------------------------
    def add_summary_table(self, module):
        self.add_caption_title(module)
        self.add_summary_table_content(module)
        self.document.add_paragraph()

    # -------------------------------------------------------------------------------------
    def add_register_table(self, module):
        for register in self.register_dict[module]['register'].keys():
            self.add_register_table_title(module, register)
            self.add_register_table_content(module, register)
            self.document.add_paragraph()

    # -------------------------------------------------------------------------------------
    def add_register_table_title(self, module, register):
        paragraph = self.add_caption_title(module)
        # paragraph.paragraph_format.left_indent = -Cm(0.25)
        paragraph.add_run(': ')
        self.add_mark_entry("OtiRegister:Name", paragraph)
        paragraph.add_run(f"{register} (")
        self.add_mark_entry(f"OtiBaseAddress:{self.register_dict[module]['docParser_base_str']}", paragraph)
        self.add_mark_entry("OtiRegister:Addr", paragraph)
        paragraph.add_run(f'0x{self.register_dict[module]["register"][register]["addr"]:04X}')
        self.add_mark_entry("OtiRegister:AddrEnd", paragraph)
        paragraph.add_run(')')
        self.add_bookmark(paragraph=paragraph, bookmark_text="", bookmark_name=register)

    # -------------------------------------------------------------------------------------
    def add_register_table_content(self, module, register):
        table = self.document.add_table(rows=1, cols=5)
        # table.left_margin  = Cm(2.25)
        shading_elm = []
        for i in range(5):  # bg in gray
            shading_elm.append(docx.oxml.parse_xml(r'<w:shd {} w:fill="d9d9d9"/>'.format(docx.oxml.ns.nsdecls('w'))))
            table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading_elm[i])

        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].width = Cm(1.0)
        hdr_cells[0].text = 'Bits'
        hdr_cells[1].text = 'Field Name'
        hdr_cells[2].text = 'Default Value'
        hdr_cells[3].width = Cm(1.0)
        hdr_cells[3].text = 'Type'
        hdr_cells[4].width = Cm(9.0)
        hdr_cells[4].text = 'Description'
        # a =  sorted(self.register_obj.get_ip_addr_field_list(ip, addr), key=lambda item: int(re.search(r"[0-9]+", str(item)).group()))

        def inser_a_row(table_obj, bits, name, default, type, description):
            row_cells = table_obj.add_row().cells

            # Bits
            row_cells[0].width = Cm(1.0)
            p = row_cells[0].paragraphs[0]
            self.add_mark_entry("OtiRegField:Bits", p)
            p.add_run(bits)

            # Name
            p = row_cells[1].paragraphs[0]
            self.add_mark_entry("OtiRegField:Name", p)
            p.add_run(name)
            self.add_mark_entry("Oti", p)

            # Default
            p = row_cells[2].paragraphs[0]
            self.add_mark_entry("OtiRegField:Default", p)
            p.add_run(default)
            self.add_mark_entry("Oti", p)

            # Type
            row_cells[3].width = Cm(1.0)
            p = row_cells[3].paragraphs[0]
            self.add_mark_entry("OtiRegField:Type", p)
            p.add_run(type)
            self.add_mark_entry("Oti", p)

            # Decription
            row_cells[4].width = Cm(9.0)
            p = row_cells[4].paragraphs[0]
            p.add_run(description)

        def range_string(high, low):  # convert [3:3] to [3]
            if high == low:
                return f"[{high}]"
            return f"[{high}:{low}]"

        field_list = list(self.register_dict[module]['register'][register]['field'].keys())

        first_field = self.register_dict[module]['register'][register]['field'][field_list[0]]
        register_high = int(first_field['high']/32+1)*32-1
        if first_field['high'] != register_high:
            inser_a_row(table, range_string(register_high, first_field['high']+1), "RSVD", "0x0", first_field["type"], "Reserved")

        for i in range(len(field_list)):
            field = self.register_dict[module]['register'][register]['field'][field_list[i]]

            inser_a_row(table, range_string(field['high'], field['low']), field_list[i], field['reset'], field['type'], field['desc'])

            if i != len(field_list)-1:
                next_field = self.register_dict[module]['register'][register]['field'][field_list[i+1]]
                if field['low']-1 != next_field['high']:
                    inser_a_row(table, range_string(field['low']-1, next_field['high']), "RSVD", "0x0", field["type"], "Reserved")

        last_field = self.register_dict[module]['register'][register]['field'][field_list[-1]]
        if last_field['low'] != 0:
            inser_a_row(table, range_string(last_field['low']-1, 0), "RSVD", "0x0", last_field["type"], "Reserved")

    # -------------------------------------------------------------------------------------
    def add_caption_title(self, register, include_chapter_nb=False):
        # Create a paragraph
        paragraph = self.document.add_paragraph('Table ', style='Caption')
        paragraph.paragraph_format.left_indent = -Cm(0.25)

        # create a caption with text
        caption_list = [' STYLEREF 1 \s', ' SEQ Table \* ARABIC'] if include_chapter_nb else [' SEQ Table \* ARABIC']
        run = paragraph.add_run()
        r = run._r
        for caption in caption_list:
            fldchar = docx.oxml.OxmlElement('w:fldChar')
            fldchar.set(docx.oxml.ns.qn('w:fldCharType'), 'begin')
            r.append(fldchar)
            instrtext = docx.oxml.OxmlElement('w:instrText')
            instrtext.text = caption
            r.append(instrtext)
            fldchar = docx.oxml.OxmlElement('w:fldChar')
            fldchar.set(docx.oxml.ns.qn('w:fldCharType'), 'end')
            r.append(fldchar)

            if include_chapter_nb:
                paragraph.add_run("-")

        paragraph.add_run(f': {register} register set')
        return paragraph
    
    # -------------------------------------------------------------------------------------
    def add_summary_table_content(self, module):
        table = self.document.add_table(rows=1, cols=3)
        shading_elm = []
        for i in range(len(table.columns)):  # put the background gray for the 1er row
            shading_elm.append(docx.oxml.parse_xml(r'<w:shd {} w:fill="d9d9d9"/>'.format(docx.oxml.ns.nsdecls('w'))))
            table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading_elm[i])

        table.style = 'Table Grid'  # add border
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Address (Hex)'
        hdr_cells[0].width = Cm(4.0)
        hdr_cells[1].text = 'Type'
        hdr_cells[1].width = Cm(1.0)
        hdr_cells[2].text = 'Name'
        hdr_cells[2].width = Cm(19.0)

        for register in self.register_dict[module]['register'].keys():
            row_cells = table.add_row().cells
            row_cells[0].text = f"0x{self.register_dict[module]['register'][register]['addr']:X}"
            row_cells[0].width = Cm(4.0)
            row_cells[1].text = self.register_dict[module]['register'][register]['type']
            row_cells[1].width = Cm(1.0)
            self.add_link(paragraph=row_cells[2].paragraphs[0], link_to=register, text=register, tool_tip=register)
            row_cells[2].width = Cm(19.0)

    # -------------------------------------------------------------------------------------
    def add_mark_entry(self, text, paragraph):
        run = paragraph.add_run()
        r = run._r
        font = run.font
        font.color.rgb = self.MARK_ENTRY_COLOR
        fldchar = docx.oxml.OxmlElement('w:fldChar')
        fldchar.set(docx.oxml.ns.qn('w:fldCharType'), 'begin')
        r.append(fldchar)

        run = paragraph.add_run()
        r = run._r
        font = run.font
        font.color.rgb = self.MARK_ENTRY_COLOR
        instrtext = docx.oxml.OxmlElement('w:instrText')
        instrtext.set(docx.oxml.ns.qn('xml:space'), 'preserve')
        instrtext.text = ' XE "%s" ' % text
        r.append(instrtext)

        run = paragraph.add_run()
        r = run._r
        font = run.font
        font.color.rgb = self.MARK_ENTRY_COLOR
        fldchar = docx.oxml.OxmlElement('w:fldChar')
        fldchar.set(docx.oxml.ns.qn('w:fldCharType'), 'end')
        r.append(fldchar)

    # -------------------------------------------------------------------------------------
    @staticmethod
    def add_bookmark(paragraph, bookmark_text, bookmark_name):
        run = paragraph.add_run()
        tag = run._r
        start = docx.oxml.OxmlElement('w:bookmarkStart')
        start.set(docx.oxml.ns.qn('w:id'), '0')
        start.set(docx.oxml.ns.qn('w:name'), bookmark_name)
        tag.append(start)

        text = docx.oxml.OxmlElement('w:r')
        text.text = bookmark_text
        tag.append(text)

        end = docx.oxml.OxmlElement('w:bookmarkEnd')
        end.set(docx.oxml.ns.qn('w:id'), '0')
        end.set(docx.oxml.ns.qn('w:name'), bookmark_name)
        tag.append(end)

    # -------------------------------------------------------------------------------------
    @staticmethod
    def add_link(paragraph, link_to, text=None, tool_tip=None):
        # create hyperlink node
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')

        # set attribute for link to bookmark
        hyperlink.set(docx.oxml.shared.qn('w:anchor'), link_to, )

        if tool_tip is not None:
            # set attribute for link to bookmark
            hyperlink.set(docx.oxml.shared.qn('w:tooltip'), tool_tip, )

        new_run = docx.oxml.shared.OxmlElement('w:r')
        new_run.append(docx.oxml.shared.OxmlElement('w:rPr'))
        new_run.text = text
        hyperlink.append(new_run)
        r = paragraph.add_run()
        r._r.append(hyperlink)
        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        
    # -------------------------------------------------------------------------------------
    # def update_fields(save_path):
    #     """ Automatically updates the fields when opening the word document """
    #     namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    #     doc = DocxTemplate(save_path)
    #
    #     element_updatefields = lxml.etree.SubElement(
    #         doc.settings.element, f"{namespace}updateFields"
    #     )
    #     element_updatefields.set(f"{namespace}val", "true")
    #
    #     doc.save(save_path)

    def set_updatefields_true(self, docx_path: str):
        """ Opens the docx and adds <w:updateFields w:val="true"/> to
           (docx_path)/word/settings.xml to enforce update of TOC (and
           other fields marked as dirty) on first open.
           Saves the file afterwards.

        Arguments:
            docx_path {str} -- Absolute path to docx
        Returns:
            Nothing
        """
        namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        doc = docx.Document(docx_path)
        # add child to doc.settings element
        element_updatefields = lxml.etree.SubElement(
            doc.settings.element, f"{namespace}updateFields"
        )
        element_updatefields.set(f"{namespace}val", "true")
        doc.save(docx_path)