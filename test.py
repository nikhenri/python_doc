# pip install docx
# pip install docxtpl

import re
for i in ["[31]", "[7]", "[31:4]", "[2:1]"]:
     x = re.search(r"[0-9]+",i)
     print(x.group())
#exit()
import json
all_project_config = None
with open("project_config.json") as f:
    all_project_config = json.load(f)

def fill_dict(dict):
    for ip_name in dict:
        print(f"working on IP {ip_name}")
        for addr in dict[ip_name]:
            if not addr.startswith("0x"):
                continue
            print(f"working on addr {addr}")
            field_found = False

            # If no field declared,
            for reg_field in dict[ip_name][addr]:  # copy parent info
                if reg_field.startswith("["):
                    field_found = True
                    break
            if not field_found:
                dict[ip_name][addr]["[31:0]"] = {}

            # copy parent info
            list_idx = []
            for reg_field in dict[ip_name][addr]:
                if reg_field.startswith("["):
                    if ":" in reg_field:
                        (high, low) = reg_field[1:-1].split(":")
                    else:
                        high = low = reg_field[1:-1]
                    dict[ip_name][addr][reg_field]["type"] = dict[ip_name][addr]["type"]
                    for i in range(int(low),int(high)+1):
                        list_idx.append(i)
                    print()
            # puts the reserved
            if len(list_idx) != 32:
                list_idx.sort()
                print(f"> {addr} doesnt descript all field")
                prev=-1
                for i in list_idx:
                    if i != prev+1:
                        print(f"missing {prev+1} to {i-1}")
                        if (i-1) - (prev+1) == 0:
                            dict[ip_name][addr][f"[{i-1}]"] ={"name":"RSVD", "rst":"0x0", "type":"RO", "desc":"Reserved"}
                        else:
                            dict[ip_name][addr][f"[{i - 1}:{prev + 1}]"] = {"name": "RSVD", "rst": "0x0", "type": "RO", "desc": "Reserved"}
                    prev = i
                if list_idx[-1] != 31:
                    print(f"missing {list_idx[-1] + 1} to 31")
                    if 31 - (list_idx[-1] + 1) == 0:
                        dict[ip_name][addr][f"[31]"] = {"name": "RSVD", "rst": "0x0", "type": "RO", "desc": "Reserved"}
                    else:
                        dict[ip_name][addr][f"[31:{list_idx[-1] + 1}]"] = {"name": "RSVD", "rst": "0x0", "type": "RO", "desc": "Reserved"}

fill_dict(all_project_config)
#exit







from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX

document = Document()
sections = document.sections
for section in sections:
    section.left_margin = Cm(2.0)


# -------------
def Table(paragraph):
    run = run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = ' STYLEREF 1 \s'
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)

    paragraph.add_run("-")

    run = run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = ' SEQ Table \* ARABIC'
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)


def MarkIndexEntry(entry,paragraph):
    run = paragraph.add_run()
    r = run._r
    font = run.font
    font.color.rgb = RGBColor(255, 192, 0)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)

    run = paragraph.add_run()
    r = run._r
    font = run.font
    font.color.rgb = RGBColor(255, 192, 0)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' XE "%s" '%(entry)
    r.append(instrText)

    run = paragraph.add_run()
    r = run._r
    font = run.font
    font.color.rgb = RGBColor(255, 192, 0)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)


def add_bookmark(paragraph, bookmark_text, bookmark_name):
    run = paragraph.add_run()
    tag = run._r
    start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
    start.set(docx.oxml.ns.qn('w:id'), '0')
    start.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    tag.append(start)

    text = docx.oxml.OxmlElement('w:r')
    text.text = bookmark_text
    tag.append(text)

    end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
    end.set(docx.oxml.ns.qn('w:id'), '0')
    end.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    tag.append(end)


def add_link(paragraph, link_to, text, tool_tip=None):
    # create hyperlink node
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')

    # set attribute for link to bookmark
    hyperlink.set(docx.oxml.shared.qn('w:anchor'), link_to, )

    if tool_tip is not None:
        # set attribute for link to bookmark
        hyperlink.set(docx.oxml.shared.qn('w:tooltip'), tool_tip, )

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    r = paragraph.add_run()
    r._r.append(hyperlink)
    #r.font.name = "Calibri"
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    #r.font.underline = True


def create_table_link(dict, ip, document):
    paragraph = document.add_paragraph('Table ', style='Caption')
    paragraph.paragraph_format.left_indent = -Cm(0.25)
    Table(paragraph)
    paragraph.add_run(f': {ip} register set')


    table = document.add_table(rows=1, cols=3)
    shading_elm = []
    for i in range(len(table.columns)):
        shading_elm.append(parse_xml(r'<w:shd {} w:fill="d9d9d9"/>'.format(nsdecls('w'))))
        table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading_elm[i])

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    #hdr_cells[0].width = Cm(1.0)
    hdr_cells[0].text = 'Address (Hex)'
    hdr_cells[0].width = Cm(4.0)
    hdr_cells[1].text = 'Type'
    hdr_cells[1].width = Cm(1.0)
    hdr_cells[2].text = 'Name'
    hdr_cells[2].width = Cm(19.0)
    #hdr_cells[3].text = 'Description'

    for addr in dict[ip]:
        if addr.startswith("0x"):
            row_cells = table.add_row().cells
            row_cells[0].text = addr
            row_cells[0].width = Cm(4.0)
            row_cells[1].text = dict[ip][addr]["type"]
            row_cells[1].width = Cm(1.0)
            # BITS
            #row_cells[0].width = Cm(1.0)
            p = row_cells[2].paragraphs[0]
            row_cells[2].width = Cm(19.0)
            #MarkIndexEntry("OtiRegField:Bits", p)
            #p.add_run(i)

            #paragraph = document.add_paragraph()
            add_link(paragraph=p, link_to=dict[ip][addr]["name"], text=dict[ip][addr]["name"], tool_tip=dict[ip][addr]["name"])

    # add a bookmark to every paragraph
    #for paranum, paragraph in enumerate(document.paragraphs):
    #    add_bookmark(paragraph=paragraph,
    #                 bookmark_text=f"", bookmark_name=f"temp{paranum + 1}")

def create_a_reg_table_caption(dict, ip, addr, document):
    paragraph = document.add_paragraph('Table ', style='Caption')
    paragraph.paragraph_format.left_indent = -Cm(0.25)
    Table(paragraph)
    paragraph.add_run(': ')
    MarkIndexEntry("OtiRegister:Name", paragraph)
    paragraph.add_run(f'{dict[ip][addr]["name"]} (')
    MarkIndexEntry("OtiBaseAddress:OTI_IPV4_BASE_ADD", paragraph)
    MarkIndexEntry("OtiRegister:Addr", paragraph)
    paragraph.add_run(f'0x{int(addr, 16):04X}')
    MarkIndexEntry("OtiRegister:AddrEnd", paragraph)
    paragraph.add_run(')')
    add_bookmark(paragraph=paragraph, bookmark_text="", bookmark_name=dict[ip][addr]["name"])

def create_a_reg_table(dict, ip, addr, document):

    table = document.add_table(rows=1, cols=5)
    #table.left_margin  = Cm(2.25)
    shading_elm = []
    for i in range(5):
        shading_elm.append(parse_xml(r'<w:shd {} w:fill="d9d9d9"/>'.format(nsdecls('w'))))
        table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading_elm[i])

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].width = Cm(1.0)
    hdr_cells[0].text = 'Bits'
    hdr_cells[1].text = 'Field Name'
    hdr_cells[2].text = 'Default Value'
    hdr_cells[3].width = Cm(1.0)
    hdr_cells[3].text = 'Type'
    hdr_cells[4].width = Cm(9.0)
    hdr_cells[4].text = 'Description'
    field_dict = {}
    for i in dict[ip][addr]:
        if i.startswith("["):
            field_dict[i] = dict[ip][addr][i]


    for i in sorted(field_dict, key=lambda item: int(re.search(r"[0-9]+", item).group()), reverse=True):
        print(f">> {i}")
        if i.startswith("["):
            row_cells = table.add_row().cells

            # BITS
            row_cells[0].width = Cm(1.0)
            p = row_cells[0].paragraphs[0]
            MarkIndexEntry("OtiRegField:Bits", p)
            p.add_run(i)

            # NAME
            p = row_cells[1].paragraphs[0]
            MarkIndexEntry("OtiRegField:Name", p)
            p.add_run(dict[ip][addr][i]["name"])
            MarkIndexEntry("Oti", p)

            # Default
            p = row_cells[2].paragraphs[0]
            MarkIndexEntry("OtiRegField:Default", p)
            p.add_run(dict[ip][addr][i]["rst"])
            MarkIndexEntry("Oti", p)

            # Type
            row_cells[3].width = Cm(1.0)
            p = row_cells[3].paragraphs[0]
            MarkIndexEntry("OtiRegField:Type", p)
            p.add_run(dict[ip][addr][i]["type"])
            MarkIndexEntry("Oti", p)


            row_cells[4].width = Cm(9.0)
            p = row_cells[4].paragraphs[0]
            p.add_run(dict[ip][addr][i]["desc"])



#create table


# creae array
for ip in all_project_config:
    create_table_link(all_project_config,ip,document)
    document.add_paragraph()
    for addr in all_project_config[ip]:
        if addr.startswith("0x"):
            create_a_reg_table_caption(all_project_config,ip,addr,document)
            create_a_reg_table(all_project_config,ip,addr,document)
            document.add_paragraph()

document.save('simple.docx')

import os
os.startfile('simple.docx')